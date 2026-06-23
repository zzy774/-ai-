#!/usr/bin/env python3
"""
Fill DOCX template with AI content. Preserves ALL formatting.
"""
import sys, os, json, re, zipfile, shutil, tempfile
from docx import Document
from docx.shared import Inches

def set_para_text(p, text):
    if not p.runs: p.add_run(text); return
    p.runs[0].text = text
    for r in p.runs[1:]: r.text = ""

def xml_replace(docx_path, replacements):
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
    with zipfile.ZipFile(docx_path) as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.namelist():
            data = zin.read(item)
            if item.endswith(".xml") or item.endswith(".rels"):
                for old, new in replacements.items():
                    data = data.replace(old.encode("utf-8"), new.encode("utf-8"))
            zout.writestr(item, data)
    shutil.move(tmp, docx_path)

def main():
    if len(sys.argv) != 4:
        print("Usage: python fill_report_template.py <template.docx> <content.json> <output.docx>", file=sys.stderr)
        sys.exit(1)

    template_path = sys.argv[1]; content_path = sys.argv[2]; output_path = sys.argv[3]

    with open(content_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    placeholders = data.get("placeholders", {})
    sections = data.get("sections", [])
    uml_images = data.get("umlImages", [])
    uml_dir = data.get("umlDir", "./uploads/uml")

    shutil.copy2(template_path, output_path)

    xml_repl = {}
    for k, v in placeholders.items():
        xml_repl["${" + k + "}"] = str(v) if v else ""
    if xml_repl:
        xml_replace(output_path, xml_repl)

    doc = Document(output_path)
    all_paras = doc.paragraphs

    section_kw = ["目的","原理","环境","步骤","结果","分析","总结","代码","设计",
                  "实现","讨论","方法","工具","要求","内容","背景","方案","实验内容",
                  "选题","需求","可行性","概要","详细","测试","部署","维护","附录",
                  "参考","致谢","引言","概述","摘要","关键词","功能","模块","架构",
                  "流程","数据","界面","接口","安全","性能","系统","设计","说明"]

    # Find ALL headings (short text with section keyword)
    all_headings = []  # (index, text)
    for i, p in enumerate(all_paras):
        text = p.text.strip()
        if not text or len(text) > 22: continue
        if "（" in text or "(" in text or "此处" in text or "..." in text or "略" in text: continue
        is_h = False
        for kw in section_kw:
            if kw in text: is_h = True; break
        if not is_h:
            for r in p.runs:
                if (r.bold and len(text) < 20) or (r.font.size and r.font.size >= 260000):
                    is_h = True; break
        if is_h:
            all_headings.append((i, text))

    # Only keep headings whose zone contains a placeholder
    full_hs = {h[0] for h in all_headings}
    target_h = []  # (index, text) — headings with placeholders
    for h_idx in range(len(all_headings)):
        start = all_headings[h_idx][0]
        end = all_headings[h_idx + 1][0] if h_idx + 1 < len(all_headings) else len(all_paras)
        for j in range(start + 1, min(start + 10, end)):
            t = all_paras[j].text.strip()
            if ('...' in t or '（略）' in t or '（此处' in t or '此处填写' in t) and len(t) < 100:
                target_h.append(all_headings[h_idx]); break

    print(f"Found {len(all_headings)} headings, {len(target_h)} with placeholders")

    if not sections:
        doc.save(output_path); print(f"OK: {output_path}"); return

    # Match+merge AI sections to template headings
    merged = {}
    for sc in sections:
        search = re.sub(r'[0-9.\s、]+', '', sc["title"]).strip()
        if len(search) < 2: continue
        best = -1
        for h in range(len(target_h)):
            htext = re.sub(r'[0-9.\s、]+', '', target_h[h][1]).strip()
            if search == htext or search in htext or htext in search:
                best = h; break
        if best < 0: continue
        if best not in merged: merged[best] = []
        merged[best].append(sc["content"])

    print(f"Merged into {len(merged)} sections")

    for best_h, contents in merged.items():
        start = target_h[best_h][0]
        end = target_h[best_h + 1][0] if best_h + 1 < len(target_h) else len(all_paras)

        # Find placeholder
        ph = None
        for i in range(start + 1, min(start + 10, end)):
            if i in full_hs: continue
            t = all_paras[i].text.strip()
            if ('...' in t or '（略）' in t or '（此处' in t or '此处填写' in t) and len(t) < 100:
                ph = all_paras[i]; break

        if ph is None: continue

        # Merge all sub-section contents
        combined = []
        for content in contents:
            full = content.strip()
            full = full.replace('**', '').replace('## ', '')
            full = re.sub(r'(?m)^[\-\*]\s+', '- ', full)
            # Fix sub-numbering from AI response header
            combined.append(full)

        set_para_text(ph, "\n\n".join(combined))
        print(f"  Filled '{target_h[best_h][1]}' ({len(''.join(combined))} chars)")

    # UML images at end
    if uml_images:
        doc.add_paragraph()
        p = doc.add_paragraph(); p.add_run("系统 UML 类图：").bold = True
        for name in uml_images:
            img_path = os.path.join(uml_dir, name)
            if os.path.exists(img_path):
                ip = doc.add_paragraph(); ip.alignment = 1
                ip.add_run().add_picture(img_path, width=Inches(4.8))
                cp = doc.add_paragraph(); cp.alignment = 1
                cp.add_run("图 " + name.replace(".png","")).italic = True

    doc.save(output_path)
    print(f"OK: {output_path}")

if __name__ == "__main__":
    main()
